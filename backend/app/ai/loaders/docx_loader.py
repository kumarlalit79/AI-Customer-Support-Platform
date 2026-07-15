from langchain_core.documents import Document
from docx import Document as DocxDocument

class DOCXLoader:

    @staticmethod
    def load(file_path: str):
        document = DocxDocument(file_path)
        text = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return [
            Document(
                page_content="\n".join(text),
                metadata={
                    "page": 0,
                    "source": file_path,
                },
            )
        ]